"""Extraction coverage for the five non-PDF formats, which had none (2026-08-20).

`extract_epub`, `extract_html`, `extract_docx`, `extract_rtf` and `extract_odt` are five of the
seven formats `SUPPORTED_EXTENSIONS` advertises, and until this file **not one line of them ran in
the suite** — `extractors.py` sat at 44% with lines 141-220 (all five functions) never executed.
Only `.txt`/`.md` and the PDF placeholder-strip were covered. Ingest is the subsystem that has
produced the most tracked defects (KI-14/26/34/40/42/43/44/46/47), and this was its front door.

**Fixtures are built here rather than committed.** Every writer used below (`python-docx`,
`odfpy`, `ebooklib`) is a base runtime dependency, so no test can skip for a missing library, and
an in-test fixture shows the reader exactly what went in next to what must come out. HTML and RTF
are hand-authored markup, so those two are producer-independent outright.

**What this deliberately does not claim.** A round-trip through the same library cannot prove the
extractor survives files from *other* producers — the real-world variance that broke tier-1
citation parsing (KI-45). These tests pin the contract, the structure heuristics and the encoding
behaviour; they are not a substitute for a real-world fixture corpus.

**Why every format asserts a non-ASCII round-trip.** Nothing on Windows defaults to UTF-8
(CONTEXT.md section 9), and this project has already shipped four double-encoded files
(`test_docs_encoding.py`). An extractor that mangles an accent turns a document into garbage that
still looks like prose, so it passes ingest, reaches the chunk store and is never noticed.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from doc_assistant.extractors import (
    extract_docx,
    extract_epub,
    extract_html,
    extract_odt,
    extract_rtf,
    extract_to_markdown,
    get_format_status,
)

# Accented Latin only: unambiguous under ruff's confusables rules (RUF001-003), unlike the dashes
# and quotes that made `test_docs_encoding.py` unable to quote its own evidence.
ACCENTED = "café résumé naïve"


# ============================================================
# Fixture builders
# ============================================================


def _write_docx(path: Path, *, empty: bool = False) -> Path:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    if not empty:
        doc.add_heading("Chapter One", level=1)
        doc.add_paragraph(f"Body prose containing {ACCENTED}.")
        doc.add_heading("A Subsection", level=2)
        doc.add_paragraph("   ")  # whitespace-only: must be dropped, not emitted as a blank part
        doc.add_heading("Deeper Still", level=3)
    doc.save(str(path))
    return path


def _write_odt(path: Path, *, empty: bool = False) -> Path:
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    doc = OpenDocumentText()
    if not empty:
        doc.text.addElement(P(text="First paragraph."))
        doc.text.addElement(P(text="   "))  # whitespace-only: must be dropped
        doc.text.addElement(P(text=f"Second paragraph with {ACCENTED}."))
    doc.save(str(path))
    return path


def _write_epub(path: Path, *, empty: bool = False) -> Path:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("id-provenote-test")
    book.set_title("A Short Treatise")
    book.set_language("en")
    chapter = epub.EpubHtml(title="Ch1", file_name="ch1.xhtml", lang="en")
    # The "empty" case is a whitespace-only paragraph, not an empty <body>: ebooklib cannot
    # serialise the latter at all (lxml raises "Document is empty" while building the nav), so a
    # literally-empty body would test the fixture builder rather than the extractor.
    body = "<p> </p>" if empty else f"<h1>Opening</h1><p>Body text containing {ACCENTED}.</p>"
    chapter.content = f"<html><body>{body}</body></html>"
    book.add_item(chapter)
    book.toc = (chapter,)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path


def _write_html(path: Path, *, empty: bool = False) -> Path:
    if empty:
        path.write_text("", encoding="utf-8")
        return path
    path.write_text(
        "<html><head><title>Doc Title</title>"
        "<style>body{color:red}</style></head>"
        "<body><nav>NAVIGATION CHROME</nav>"
        "<h1>Main Heading</h1>"
        f"<p>Body sentence containing {ACCENTED}.</p>"
        "<h2>Sub Heading</h2><p>More prose.</p>"
        "<script>var tracking = 1;</script>"
        "<footer>FOOTER CHROME</footer></body></html>",
        encoding="utf-8",
    )
    return path


def _write_rtf(path: Path, *, empty: bool = False) -> Path:
    if empty:
        path.write_text(r"{\rtf1\ansi\deff0 }", encoding="utf-8")
        return path
    # \'e9 / \'ef are the ANSI escapes RTF uses for non-ASCII; striprtf must decode them back.
    path.write_text(
        r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Times;}}\f0\fs24 "
        r"Opening line.\par Body line with caf\'e9 r\'e9sum\'e9 na\'efve.\par}",
        encoding="utf-8",
    )
    return path


#: (suffix, builder) for the dispatch + robustness tests that treat all five uniformly.
FORMATS = [
    (".docx", _write_docx),
    (".odt", _write_odt),
    (".epub", _write_epub),
    (".html", _write_html),
    (".rtf", _write_rtf),
]


# ============================================================
# DOCX
# ============================================================


def test_docx_keeps_body_prose(tmp_path: Path) -> None:
    md = extract_docx(_write_docx(tmp_path / "a.docx"))
    assert "Body prose containing" in md


def test_docx_maps_heading_styles_to_markdown_levels(tmp_path: Path) -> None:
    """The style-name heuristic in `extract_docx` is the only structure DOCX carries through."""
    md = extract_docx(_write_docx(tmp_path / "a.docx"))
    assert "# Chapter One" in md
    assert "## A Subsection" in md
    assert "### Deeper Still" in md


def test_docx_drops_whitespace_only_paragraphs(tmp_path: Path) -> None:
    """A blank paragraph must not become an empty part — it would widen every join by one gap."""
    md = extract_docx(_write_docx(tmp_path / "a.docx"))
    assert "\n\n\n" not in md


def test_docx_round_trips_non_ascii(tmp_path: Path) -> None:
    assert ACCENTED in extract_docx(_write_docx(tmp_path / "a.docx"))


# ============================================================
# ODT
# ============================================================


def test_odt_keeps_every_paragraph_in_order(tmp_path: Path) -> None:
    md = extract_odt(_write_odt(tmp_path / "a.odt"))
    assert md.index("First paragraph.") < md.index("Second paragraph")


def test_odt_drops_whitespace_only_paragraphs(tmp_path: Path) -> None:
    assert "\n\n\n" not in extract_odt(_write_odt(tmp_path / "a.odt"))


def test_odt_round_trips_non_ascii(tmp_path: Path) -> None:
    assert ACCENTED in extract_odt(_write_odt(tmp_path / "a.odt"))


# ============================================================
# EPUB
# ============================================================


def test_epub_promotes_the_dublin_core_title_to_an_h1(tmp_path: Path) -> None:
    """The DC title is the only metadata the extractor lifts, and it leads the markdown."""
    md = extract_epub(_write_epub(tmp_path / "a.epub"))
    assert md.startswith("# A Short Treatise")


def test_epub_keeps_inner_html_headings_and_prose(tmp_path: Path) -> None:
    md = extract_epub(_write_epub(tmp_path / "a.epub"))
    assert "# Opening" in md
    assert "Body text containing" in md


def test_epub_round_trips_non_ascii(tmp_path: Path) -> None:
    assert ACCENTED in extract_epub(_write_epub(tmp_path / "a.epub"))


# ============================================================
# HTML
# ============================================================


@pytest.mark.parametrize(
    ("chrome", "marker"),
    [
        ("nav", "NAVIGATION CHROME"),
        ("footer", "FOOTER CHROME"),
        ("script", "var tracking"),
        ("style", "color:red"),
    ],
)
def test_html_discards_page_chrome(tmp_path: Path, chrome: str, marker: str) -> None:
    """`extract_html` decomposes script/style/nav/footer — the one content decision it makes.

    Stated per tag so a future edit to that tuple fails on the tag it dropped, not on a single
    opaque assertion covering all four.
    """
    md = extract_html(_write_html(tmp_path / "a.html"))
    assert marker not in md, f"<{chrome}> content survived extraction"


def test_html_converts_headings_and_keeps_prose(tmp_path: Path) -> None:
    md = extract_html(_write_html(tmp_path / "a.html"))
    assert "# Main Heading" in md
    assert "## Sub Heading" in md
    assert "More prose." in md


def test_html_round_trips_non_ascii(tmp_path: Path) -> None:
    assert ACCENTED in extract_html(_write_html(tmp_path / "a.html"))


def test_html_reads_as_utf8_regardless_of_the_ansi_codepage(tmp_path: Path) -> None:
    """Pins the explicit `encoding="utf-8"` on the read (CONTEXT.md section 9).

    Without it Python would use the ANSI codepage on Windows and this file's accents would arrive
    as mojibake, which reads as prose and so survives every downstream check.
    """
    path = tmp_path / "a.html"
    path.write_bytes(f"<html><body><p>{ACCENTED}</p></body></html>".encode())
    assert ACCENTED in extract_html(path)


# ============================================================
# RTF
# ============================================================


def test_rtf_strips_control_words_and_keeps_the_text(tmp_path: Path) -> None:
    md = extract_rtf(_write_rtf(tmp_path / "a.rtf"))
    assert "Opening line." in md
    assert "rtf1" not in md
    assert "fonttbl" not in md


def test_rtf_decodes_ansi_escapes_back_to_characters(tmp_path: Path) -> None:
    r"""RTF stores non-ASCII as `\'e9`-style escapes; left literal they would poison the text."""
    md = extract_rtf(_write_rtf(tmp_path / "a.rtf"))
    assert ACCENTED in md
    assert "\\'e9" not in md


# ============================================================
# Dispatch + robustness, stated once over all five
# ============================================================


@pytest.mark.parametrize(("suffix", "builder"), FORMATS)
def test_extract_to_markdown_dispatches_every_supported_format(
    tmp_path: Path, suffix: str, builder: object
) -> None:
    """The dispatch table is the real entry point: a format wired nowhere is gone."""
    path = builder(tmp_path / f"sample{suffix}")  # type: ignore[operator]
    assert extract_to_markdown(path).strip()


@pytest.mark.parametrize(("suffix", "builder"), FORMATS)
def test_a_valid_but_empty_document_returns_empty_text_not_a_crash(
    tmp_path: Path, suffix: str, builder: object
) -> None:
    """Degrade honestly at zero content (robustness contract) rather than raising at ingest."""
    path = builder(tmp_path / f"empty{suffix}", empty=True)  # type: ignore[operator]
    assert isinstance(extract_to_markdown(path), str)


def test_htm_and_html_reach_the_same_extractor(tmp_path: Path) -> None:
    """`.htm` is a separate key in the dispatch table and is easy to drop when editing it."""
    body = "<html><body><p>Shared body.</p></body></html>"
    (tmp_path / "a.htm").write_text(body, encoding="utf-8")
    (tmp_path / "a.html").write_text(body, encoding="utf-8")
    assert extract_to_markdown(tmp_path / "a.htm") == extract_to_markdown(tmp_path / "a.html")


def test_suffix_matching_is_case_insensitive_through_the_dispatch(tmp_path: Path) -> None:
    path = _write_html(tmp_path / "a.HTML")
    assert "# Main Heading" in extract_to_markdown(path)


def test_a_non_pymupdf_pdf_extractor_is_refused_by_name(tmp_path: Path) -> None:
    """Marker was removed from the production path; the refusal must say so rather than fall back.

    The guard fires before any parsing, so this needs no PDF — which is the point: silently using
    a different extractor would change every downstream hash without anything reporting it.
    """
    path = tmp_path / "paper.pdf"
    path.write_bytes(b"%PDF-1.4 not a real document")
    with pytest.raises(ValueError, match="marker"):
        extract_to_markdown(path, pdf_extractor="marker")


# ============================================================
# get_format_status — the advisory the UI shows for a rejected file
# ============================================================


def test_get_format_status_accepts_supported_formats_without_advice() -> None:
    assert get_format_status(Path("paper.pdf")) == (True, None)
    assert get_format_status(Path("book.epub")) == (True, None)


@pytest.mark.parametrize(
    ("name", "hint"),
    [("old.doc", "DOCX"), ("paper.tex", "PDF"), ("book.mobi", "EPUB")],
)
def test_get_format_status_names_the_conversion_target(name: str, hint: str) -> None:
    """A rejection that does not say what to do instead is a dead end for the user."""
    supported, advice = get_format_status(Path(name))
    assert not supported
    assert advice is not None
    assert hint in advice


def test_get_format_status_still_advises_on_an_unmapped_extension() -> None:
    supported, advice = get_format_status(Path("photo.jpg"))
    assert not supported
    assert advice is not None
    assert ".jpg" in advice


# ============================================================
# `_soup_to_markdown` — the rule EPUB and HTML now share (2026-08-20)
# ============================================================
#
# Unifying the two extractors onto one helper gave EPUB three behaviours it never had: chrome
# removal, `<head>` removal, and inline-tag unwrapping. Those are asserted here rather than in
# `test_extraction_fixtures.py`, because reaching them needs chapter markup the committed fixture
# deliberately does not carry.


def _epub_with_chapter_markup(path: Path, body: str) -> Path:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("id-chapter-markup")
    book.set_title("Book Title")
    book.set_language("en")
    chapter = epub.EpubHtml(title="Ch1", file_name="ch1.xhtml", lang="en")
    chapter.content = (
        f"<html><head><title>CHAPTER HEAD TITLE</title></head><body>{body}</body></html>"
    )
    book.add_item(chapter)
    book.toc = (chapter,)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path


def test_epub_chapter_head_titles_do_not_reach_the_text(tmp_path: Path) -> None:
    """Real publisher EPUBs put a `<title>` in every chapter's `<head>`; before the shared helper
    each one landed in the body as a bare line."""
    path = _epub_with_chapter_markup(tmp_path / "a.epub", "<p>Real body prose.</p>")
    md = extract_epub(path)
    assert "Real body prose." in md
    assert "CHAPTER HEAD TITLE" not in md


def test_epub_drops_page_chrome_the_same_way_html_does(tmp_path: Path) -> None:
    """EPUB never removed script/style/nav/footer — the two extractors held the same rule twice
    and had drifted. One helper now, so they cannot disagree again."""
    path = _epub_with_chapter_markup(
        tmp_path / "a.epub",
        "<nav>NAV CHROME</nav><p>Body.</p><script>SCRIPT CHROME</script>"
        "<style>STYLE CHROME</style><footer>FOOTER CHROME</footer>",
    )
    md = extract_epub(path)
    assert "Body." in md
    for marker in ("NAV CHROME", "SCRIPT CHROME", "STYLE CHROME", "FOOTER CHROME"):
        assert marker not in md, f"{marker} survived EPUB extraction"


def test_epub_keeps_a_sentence_whole_across_inline_markup(tmp_path: Path) -> None:
    path = _epub_with_chapter_markup(
        tmp_path / "a.epub", "<p>The <em>Drosophila</em> gene <strong>white</strong> is named.</p>"
    )
    lines = extract_epub(path).splitlines()
    assert any("Drosophila" in line and "white" in line for line in lines)


def test_block_elements_still_separate_after_inline_unwrapping(tmp_path: Path) -> None:
    """The guard on the fix: `separator="\n"` is what keeps paragraphs and list items apart, so
    unwrapping inline tags must not let two blocks run together into one line."""
    path = _write_html(tmp_path / "a.html")
    lines = [line for line in extract_html(path).splitlines() if line.strip()]
    assert any(line.strip() == "More prose." for line in lines), (
        "a block element merged into its neighbour"
    )
